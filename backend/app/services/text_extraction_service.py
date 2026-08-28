from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader
import fitz  # PyMuPDF

from app.services.ocr_service import OCRService


class TextExtractionService:
    """
    Extract text from supported contract files.

    PDF extraction uses PyMuPDF first because it preserves
    the visual reading order of PDF content better than pypdf,
    especially for tables such as Schedule A.

    Falls back to pypdf and then OCR when necessary.
    """

    @staticmethod
    def extract(
        filename: str,
        content: bytes,
    ) -> str:

        extension = Path(filename).suffix.lower()

        if extension == ".txt":
            return TextExtractionService._extract_txt(content)

        if extension == ".pdf":
            return TextExtractionService._extract_pdf(content)

        if extension == ".docx":
            return TextExtractionService._extract_docx(content)

        if extension in {".png", ".jpg", ".jpeg"}:
            return OCRService.extract_from_image(content)

        raise ValueError(
            f"Unsupported file type for text extraction: "
            f"{extension or 'unknown'}"
        )

    # ================================================================
    # TXT
    # ================================================================

    @staticmethod
    def _extract_txt(content: bytes) -> str:

        try:
            text = content.decode("utf-8")

        except UnicodeDecodeError:
            text = content.decode(
                "utf-8",
                errors="replace",
            )

        return text.strip()

    # ================================================================
    # PDF
    # ================================================================

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        """
        Extract PDF text using PyMuPDF.

        We deliberately use blocks instead of plain page.get_text()
        because blocks retain more layout information.

        This is especially useful for:
        - Schedule A
        - loan particulars
        - amount fields
        - interest fields
        - repayment fields
        - fee/charge tables
        """

        # ------------------------------------------------------------
        # 1. PyMuPDF
        # ------------------------------------------------------------

        try:

            document = fitz.open(
                stream=content,
                filetype="pdf",
            )

            pages = []

            for page in document:

                blocks = page.get_text(
                    "blocks",
                    sort=True,
                )

                page_parts = []

                for block in blocks:

                    if len(block) < 5:
                        continue

                    block_text = block[4]

                    if not block_text:
                        continue

                    block_text = (
                        block_text
                        .replace("\x00", "")
                        .strip()
                    )

                    if block_text:
                        page_parts.append(
                            block_text
                        )

                if page_parts:
                    pages.append(
                        "\n".join(page_parts)
                    )

            document.close()

            extracted_text = (
                "\n\n".join(pages)
                .strip()
            )

            if extracted_text:
                return extracted_text

        except Exception:
            # Continue to pypdf fallback.
            pass

        # ------------------------------------------------------------
        # 2. pypdf fallback
        # ------------------------------------------------------------

        try:

            reader = PdfReader(
                BytesIO(content)
            )

            pages = []

            for page in reader.pages:

                page_text = page.extract_text()

                if page_text:
                    pages.append(
                        page_text
                    )

            extracted_text = (
                "\n\n".join(pages)
                .strip()
            )

            if extracted_text:
                return extracted_text

        except Exception:
            pass

        # ------------------------------------------------------------
        # 3. OCR fallback
        # ------------------------------------------------------------

        return OCRService.extract_from_pdf(
            content
        )

    # ================================================================
    # DOCX
    # ================================================================

    @staticmethod
    def _extract_docx(content: bytes) -> str:

        document = Document(
            BytesIO(content)
        )

        paragraphs = []

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(
                    text
                )

        # ------------------------------------------------------------
        # Also extract DOCX table content.
        # ------------------------------------------------------------

        tables = []

        for table in document.tables:

            rows = []

            for row in table.rows:

                cells = []

                for cell in row.cells:

                    cell_text = (
                        cell.text
                        .replace("\n", " ")
                        .strip()
                    )

                    cells.append(
                        cell_text
                    )

                if any(cells):
                    rows.append(
                        " | ".join(cells)
                    )

            if rows:
                tables.append(
                    "\n".join(rows)
                )

        parts = []

        if paragraphs:
            parts.append(
                "\n\n".join(paragraphs)
            )

        if tables:
            parts.append(
                "\n\n".join(tables)
            )

        return "\n\n".join(parts).strip()